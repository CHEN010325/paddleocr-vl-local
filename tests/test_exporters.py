import base64
import io
from pathlib import Path

import pytest
from lxml import etree
from PIL import Image
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas

from exporters import (
    ExportError,
    _image_bytes,
    _load_formula_rasterizer,
    _formula_image,
    extract_html_table_specs,
    _invisible_overlay,
    _formula_omml,
    _normalize_latex_source,
    _parse_html_image_line,
    _page_number,
    _preferred_cjk_font,
    _result_text_lines,
    _split_formula_cells,
    _split_formula_rows,
    _task_page_text,
    build_docx,
    build_searchable_pdf,
    build_xlsx,
    extract_html_tables,
    extract_markdown_tables,
    plain_markdown_text,
    split_markdown_table_row,
)


SAMPLE_MARKDOWN = """# 项目报告

这是包含 **重点**、`code` 和 [链接](https://example.com) 的正文。

- 第一项
- 第二项

| 名称 | 数值 | 备注 |
| --- | ---: | :--- |
| Alpha | 12 | 正常 |
| Formula | =2+2 | 防注入 |

> 本地离线处理。
"""


def sample_task(**updates):
    task = {
        "name": "示例报告.pdf",
        "modelName": "PP-OCRv6",
        "mimeType": "application/pdf",
        "markdown": SAMPLE_MARKDOWN,
        "images": {},
        "ocrResults": [
            {"sourcePage": 1, "ocrLines": [{"text": "Searchable first page"}, {"text": "中文第一页"}]},
            {"sourcePage": 2, "markdown": {"text": "Searchable second page"}},
        ],
    }
    task.update(updates)
    return task


def make_pdf(path: Path, pages: int = 2):
    pdf = canvas.Canvas(str(path), pagesize=(300, 420))
    for index in range(pages):
        pdf.drawString(30, 390, f"Original page {index + 1}")
        pdf.showPage()
    pdf.save()


def test_markdown_table_parser_handles_escaped_cells_and_fences():
    markdown = """| A | B |
| --- | --- |
| x\\|y | z |

```md
| fake | table |
| --- | --- |
```
"""
    assert split_markdown_table_row(r"| x\|y | z |") == ["x|y", "z"]
    assert extract_markdown_tables(markdown) == [[["A", "B"], ["x|y", "z"]]]


def test_markdown_table_parser_preserves_backslashes_and_pipe_escape_parity():
    assert split_markdown_table_row(r"| Path | C:\temp\file.txt |") == ["Path", r"C:\temp\file.txt"]
    assert split_markdown_table_row(r"| Formula | \alpha + \beta |") == ["Formula", r"\alpha + \beta"]
    assert split_markdown_table_row(r"| Escaped | x\|y |") == ["Escaped", "x|y"]
    assert split_markdown_table_row(r"| Even | left\\|right |") == ["Even", r"left\\", "right"]
    assert split_markdown_table_row(r"| Odd | left\\\|right | tail |") == [
        "Odd",
        r"left\\|right",
        "tail",
    ]
    assert split_markdown_table_row(r"| Tail | value\|") == ["Tail", "value|"]


def test_plain_markdown_text_preserves_literal_symbols_and_unwraps_real_markup():
    literal = r"foo_bar_baz foo__bar__baz x < y > z C:\temp\file.txt \alpha + \beta"
    assert plain_markdown_text(literal) == literal
    assert plain_markdown_text(r"x\|y") == "x|y"
    assert plain_markdown_text(r"left\\|right") == r"left\\|right"
    assert plain_markdown_text("**bold** __strong__ *italic* _emphasis_ `code`") == (
        "bold strong italic emphasis code"
    )
    assert plain_markdown_text("<strong>bold</strong> <em>italic</em><br>next") == "bold italic\nnext"


def test_plain_markdown_text_preserves_currency_and_escaped_dollars():
    # A pair of currency amounts must not be consumed as one giant formula.
    assert plain_markdown_text("Cost is $5 and tax is $10.") == "Cost is $5 and tax is $10."
    assert plain_markdown_text("Price $5.00 (USD), then text.") == "Price $5.00 (USD), then text."
    assert plain_markdown_text("$ (r) $") == "(r)"
    # Markdown-escaped dollars are literal text, even when another valid
    # formula follows later in the same paragraph.
    assert plain_markdown_text(r"Literal \$x\$ and \$5; formula $y^2$") == (
        "Literal $x$ and $5; formula y^2"
    )
    assert plain_markdown_text(r"Literal \$x then $y$") == "Literal $x then y"


def test_bare_equation_and_align_environments_are_exported_as_math():
    markdown = r"""Bare equation: \begin{equation}x^2 + y^2\tag{1}\end{equation}

\begin{align}a&=b\\c&=d\tag{2}\end{align}

\begin{array}{ll}u&=v\\w&=z\end{array}
"""
    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is not None:
        assert len(document.inline_shapes) >= 3
    # Environment wrappers are not leaked as literal OCR text into the DOCX.
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "\\begin" not in paragraph_text
    assert "\\end" not in paragraph_text
    assert "(1)" in paragraph_text
    assert "(2)" in paragraph_text


def test_legacy_arrayll_shorthand_is_detected_without_math_delimiters():
    markdown = r"""Equation: (r,a)=\\arrayll R(q,r,a) \\& iteration=0 \\\\ R(q,d_k^1,r,a) \\& otherwise \\array (5)
"""
    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is not None:
        assert len(document.inline_shapes) == 1
    with io.BytesIO(build_docx(sample_task(markdown=markdown))) as stream:
        import zipfile

        with zipfile.ZipFile(stream) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
    assert "arrayll" not in xml
    assert "R(q,r,a)" not in "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "(5)" in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_legacy_arrayll_double_escaping_keeps_rows_and_subscripts_aligned():
    # Some JSON/Markdown layers double every backslash and escape alignment
    # ampersands.  The renderer must restore the two-column array instead of
    # drawing a literal ``\\``/``&`` string or losing the subscript.
    markdown = r"""(r,a)=\\arrayll LLM.generate(q)\\&if iteration=0\\\\ LLM.generate(q,\\d_k^1)\\&otherwise\\array (5)"""
    match = next(__import__("exporters")._iter_formula_matches(markdown))
    source = __import__("exporters")._formula_source_from_match(match)
    normalized = __import__("exporters")._normalize_latex_source(source)
    body = __import__("exporters").ARRAY_ENV_RE.search(normalized).group("body")
    rows = _split_formula_rows(body)
    assert [_split_formula_cells(row) for row in rows] == [
        ["LLM.generate(q)", "if iteration=0"],
        ["LLM.generate(q,d_k^1)", "otherwise"],
    ]
    payload = build_docx(sample_task(markdown=markdown))
    document = Document(io.BytesIO(payload))
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is not None:
        assert len(document.inline_shapes) == 1


def test_html_table_br_is_preserved_in_docx_and_xlsx():
    markdown = (
        "<table><tr><th>Method</th><th>Details</th></tr>"
        "<tr><td>A</td><td>line one<br>line two</td></tr></table>"
    )
    spec = extract_html_table_specs(markdown)[0]
    assert spec.rows[1][1] == "line one\nline two"
    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    assert document.tables[0].cell(1, 1).text == "line one\nline two"
    workbook = load_workbook(io.BytesIO(build_xlsx(sample_task(markdown=markdown))))
    assert workbook["Table 1"]["B2"].value == "line one\nline two"


def test_markdown_helpers_handle_boundary_inputs(monkeypatch):
    monkeypatch.setenv("PANDOCR_DOCX_CJK_FONT", "Project Test Font")
    assert _preferred_cjk_font() == "Project Test Font"
    assert split_markdown_table_row("a|b\\") == ["a", "b\\"]
    assert split_markdown_table_row("| a | b") == ["a", "b"]
    assert extract_markdown_tables("single cell\n---\n") == []
    assert extract_markdown_tables("| a | b |\n| --- | --- |\ninvalid\n") == [[["a", "b"]]]


def test_docx_export_is_editable_and_uses_expected_structure():
    payload = build_docx(sample_task())
    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "示例报告" in text
    assert "项目报告" in text
    assert "第一项" in text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == "Alpha"
    assert document.sections[0].top_margin.inches == pytest.approx(1.0, abs=0.01)
    assert document.styles["Normal"].paragraph_format.line_spacing == pytest.approx(1.25)


def test_docx_export_handles_images_code_numbered_lists_and_fallbacks():
    image = io.BytesIO()
    Image.new("RGB", (120, 60), "navy").save(image, format="PNG")
    encoded = base64.b64encode(image.getvalue()).decode()
    markdown = """## Rich content

1. Numbered *item*

![Embedded preview](asset.png)

![Unavailable preview](missing.png)

---

```python
print('offline')
```
"""
    task = sample_task(markdown=markdown, images={"asset.png": f"data:image/png;base64,{encoded}"})
    payload = build_docx(task)
    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Numbered item" in text
    assert "Embedded preview" in text
    assert "Unavailable preview" in text
    assert "print('offline')" in text
    assert len(document.inline_shapes) == 1
    assert any(paragraph.style.name == "List Number" for paragraph in document.paragraphs)


def test_docx_export_resolves_paddleocr_vl_html_images_and_tag_star_formulas():
    image = io.BytesIO()
    Image.new("RGB", (160, 90), "navy").save(image, format="JPEG")
    encoded = base64.b64encode(image.getvalue()).decode()
    markdown = (
        '<div style="text-align: center;"><img src="imgs/chart.jpg" alt="Chart" width="36%" /></div>\n\n'
        '$$\\mathrm{m o d e l}=\\operatorname{N o r m}(x)\\tag*{(1)}$$'
    )
    task = sample_task(
        markdown=markdown,
        images={"ocr_images/batch_chart.jpg": encoded},
        ocrResults=[
            {
                "markdown": {
                    "images": {"imgs/chart.jpg": "ocr_images/batch_chart.jpg"}
                }
            }
        ],
    )
    assert _parse_html_image_line(markdown.splitlines()[0]) == ("Chart", "imgs/chart.jpg", pytest.approx(2.25))
    assert __import__("exporters")._formula_parts(r"\mathrm{m o d e l}\tag*{(1)}") == (
        r"\mathrm{model}",
        "1",
    )
    document = Document(io.BytesIO(build_docx(task)))
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is not None:
        assert len(document.inline_shapes) == 2
    assert "Chart" in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_docx_export_handles_unclosed_fence_and_requires_markdown():
    payload = build_docx(sample_task(markdown="Before\n\n```\nunclosed code"))
    document = Document(io.BytesIO(payload))
    assert "unclosed code" in "\n".join(paragraph.text for paragraph in document.paragraphs)
    with pytest.raises(ExportError, match="no Markdown result"):
        build_docx(sample_task(markdown=""))


def test_docx_and_xlsx_exports_preserve_literal_special_text():
    markdown = r"""# Literal text

foo_bar_baz foo__bar__baz

x < y > z

C:\temp\file.txt and \alpha + \beta

**bold** _italic_ <strong>HTML bold</strong>

| Kind | Value |
| --- | --- |
| Path | C:\temp\file.txt |
| Formula | \alpha + \beta |
| Symbols | foo_bar_baz and x < y > z |
| Escaped pipe | x\|y |
"""

    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "foo_bar_baz foo__bar__baz" in paragraph_text
    assert "x < y > z" in paragraph_text
    assert r"C:\temp\file.txt and \alpha + \beta" in paragraph_text
    assert "bold italic HTML bold" in paragraph_text
    assert document.tables[0].cell(1, 1).text == r"C:\temp\file.txt"
    assert document.tables[0].cell(2, 1).text == r"\alpha + \beta"
    assert document.tables[0].cell(3, 1).text == "foo_bar_baz and x < y > z"
    assert document.tables[0].cell(4, 1).text == "x|y"

    workbook = load_workbook(io.BytesIO(build_xlsx(sample_task(markdown=markdown))))
    sheet = workbook["Table 1"]
    assert sheet["B2"].value == r"C:\temp\file.txt"
    assert sheet["B3"].value == r"\alpha + \beta"
    assert sheet["B4"].value == "foo_bar_baz and x < y > z"
    assert sheet["B5"].value == "x|y"


def test_image_payload_validation_rejects_unknown_invalid_and_oversized_data(monkeypatch):
    task = {"images": {"invalid": "not-base64", "too-big": "payload"}}
    assert _image_bytes(task, "missing") is None
    assert _image_bytes(task, "invalid") is None
    monkeypatch.setattr("exporters.base64.b64decode", lambda *_args, **_kwargs: b"x" * (25 * 1024 * 1024 + 1))
    assert _image_bytes(task, "too-big") is None


def test_xlsx_export_preserves_tables_and_neutralizes_formulas():
    payload = build_xlsx(sample_task())
    workbook = load_workbook(io.BytesIO(payload), data_only=False)
    sheet = workbook["Table 1"]
    assert sheet["A1"].value == "名称"
    assert sheet["B2"].value == "12"
    assert sheet["B3"].value == "'=2+2"
    assert sheet.freeze_panes == "A2"
    assert sheet.auto_filter.ref == "A1:C3"


def test_xlsx_export_requires_at_least_one_table():
    with pytest.raises(ExportError, match="no Markdown tables"):
        build_xlsx(sample_task(markdown="Plain text"))


def test_xlsx_export_supports_multiple_tables_and_caps_column_width():
    markdown = f"| A | B |\n| --- | --- |\n| {'x' * 100} | @cmd |\n\n| C | D |\n| --- | --- |\n| +1 | -2 |"
    workbook = load_workbook(io.BytesIO(build_xlsx(sample_task(markdown=markdown))))
    assert workbook.sheetnames == ["Table 1", "Table 2"]
    assert workbook["Table 1"].column_dimensions["A"].width == 60
    assert workbook["Table 1"]["B2"].value == "'@cmd"
    assert workbook["Table 2"]["A2"].value == "'+1"
    assert workbook["Table 2"]["B2"].value == "'-2"


def test_searchable_pdf_keeps_pages_and_adds_text_layer(tmp_path):
    source = tmp_path / "source.pdf"
    make_pdf(source)
    payload = build_searchable_pdf(sample_task(), source)
    reader = PdfReader(io.BytesIO(payload))
    assert len(reader.pages) == 2
    assert "Original page 1" in (reader.pages[0].extract_text() or "")
    assert "Searchable first page" in (reader.pages[0].extract_text() or "")
    assert "Searchable second page" in (reader.pages[1].extract_text() or "")
    assert reader.metadata.producer == "PaddleOCR Local"


def test_searchable_pdf_uses_manually_edited_paginated_markdown(tmp_path):
    source = tmp_path / "source.pdf"
    make_pdf(source)
    task = sample_task(
        manualEditedAt=1,
        markdown="Corrected first page\n\n---\n\nCorrected second page",
        ocrResults=[
            {"sourcePage": 1, "ocrLines": [{"text": "STALE OCR FIRST"}]},
            {"sourcePage": 2, "ocrLines": [{"text": "STALE OCR SECOND"}]},
        ],
    )

    assert _task_page_text(task, 2) == [["Corrected first page"], ["Corrected second page"]]
    reader = PdfReader(io.BytesIO(build_searchable_pdf(task, source)))
    first_page = reader.pages[0].extract_text() or ""
    second_page = reader.pages[1].extract_text() or ""
    assert "Corrected first page" in first_page
    assert "Corrected second page" in second_page
    assert "STALE OCR FIRST" not in first_page
    assert "STALE OCR SECOND" not in second_page


def test_searchable_pdf_unescapes_literal_pipe_in_text_layer(tmp_path):
    source = tmp_path / "source.pdf"
    make_pdf(source, pages=1)
    task = sample_task(
        manualEditedAt=1,
        markdown=r"Literal x\|y and C:\temp\file.txt",
        ocrResults=[],
    )
    extracted = PdfReader(io.BytesIO(build_searchable_pdf(task, source))).pages[0].extract_text() or ""
    assert "Literal x|y" in extracted
    assert r"x\|y" not in extracted
    assert r"C:\temp\file.txt" in extracted


def test_searchable_pdf_accepts_an_image_source(tmp_path):
    source = tmp_path / "source.png"
    Image.new("RGB", (320, 180), "white").save(source)
    task = sample_task(mimeType="image/png", ocrResults=[{"sourcePage": 1, "text": "Image OCR text"}])
    reader = PdfReader(io.BytesIO(build_searchable_pdf(task, source)))
    assert len(reader.pages) == 1
    assert "Image OCR text" in (reader.pages[0].extract_text() or "")


def test_searchable_pdf_validates_source_and_ocr_text(tmp_path):
    missing = tmp_path / "missing.pdf"
    with pytest.raises(ExportError, match="source file is missing"):
        build_searchable_pdf(sample_task(), missing)

    invalid = tmp_path / "invalid.pdf"
    invalid.write_bytes(b"not a PDF")
    with pytest.raises(ExportError, match="invalid or encrypted"):
        build_searchable_pdf(sample_task(), invalid)

    unsupported = tmp_path / "source.bin"
    unsupported.write_bytes(b"not an image")
    with pytest.raises(ExportError, match="not a supported PDF or image"):
        build_searchable_pdf(sample_task(mimeType="application/octet-stream"), unsupported)

    source = tmp_path / "empty.pdf"
    make_pdf(source, pages=1)
    with pytest.raises(ExportError, match="no OCR text"):
        build_searchable_pdf(sample_task(markdown="", ocrResults=[]), source)


def test_searchable_pdf_rejects_password_encrypted_source_with_export_error(tmp_path):
    source = tmp_path / "source.pdf"
    make_pdf(source, pages=1)
    reader = PdfReader(source)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("secret")
    encrypted = tmp_path / "encrypted.pdf"
    with encrypted.open("wb") as stream:
        writer.write(stream)

    with pytest.raises(ExportError, match="encrypted"):
        build_searchable_pdf(sample_task(), encrypted)


def test_page_text_helpers_map_formats_deduplicate_and_fallback_to_markdown():
    assert _page_number({"sourcePage": 0}, 3) == 1
    assert _page_number({"pageIndex": 1.9}, 3) == 2
    assert _page_number({"page_index": 4}, 3) == 5
    assert _page_number({"sourcePage": True}, 3) == 3
    assert _result_text_lines(
        {
            "ocrLines": [{"text": "same"}, "ignored"],
            "parsing_res_list": [{"block_content": "same\nnew"}, "ignored"],
            "text": "unused",
        }
    ) == ["same", "new"]
    assert _result_text_lines({"parsingResList": [{"blockContent": "block"}]}) == ["block"]
    assert _result_text_lines({"markdown": {"text": "md line"}}) == ["md line"]
    assert _result_text_lines({"text": "plain line"}) == ["plain line"]

    task = {"ocrResults": ["ignored"], "markdown": "first\n\n---\n\nsecond"}
    assert _task_page_text(task, 2) == [["first"], ["second"]]
    assert _task_page_text({"markdown": "single\npage"}, 2) == [["single", "page"], []]
    assert _task_page_text({"markdown": ""}, 1) == [[]]


def test_invisible_overlay_wraps_long_text_and_remains_extractable():
    long_text = "Searchable " + "x" * 1000 + "\x00 control"
    reader = PdfReader(io.BytesIO(_invisible_overlay((72, 72), [long_text])))
    extracted = reader.pages[0].extract_text() or ""
    assert "Searchable" in extracted
    assert "control" in extracted


def test_html_tables_remain_structured_in_docx_and_xlsx():
    markdown = "Before\n\n<table><tr><th>Methods</th><th>ACC</th></tr><tr><td>FLARE</td><td>0.82</td></tr></table>\n\nAfter"
    assert extract_html_tables(markdown) == [[['Methods', 'ACC'], ['FLARE', '0.82']]]
    task = sample_task(markdown=markdown)
    document = Document(io.BytesIO(build_docx(task)))
    assert len(document.tables) == 1
    assert [[cell.text for cell in row.cells] for row in document.tables[0].rows] == [
        ['Methods', 'ACC'], ['FLARE', '0.82']
    ]
    workbook = load_workbook(io.BytesIO(build_xlsx(task)))
    assert list(workbook['Table 1'].values) == [('Methods', 'ACC'), ('FLARE', '0.82')]


def test_html_table_spans_are_expanded_and_merged_in_docx():
    markdown = (
        '<table><tr><td rowspan="2">Method</td><td colspan="2">Score</td></tr>'
        '<tr><td>EM</td><td>ACC</td></tr><tr><td>Demo</td><td>1</td><td>2</td></tr></table>'
    )
    spec = extract_html_table_specs(markdown)[0]
    assert len(spec.rows) == 3
    assert len(spec.rows[0]) == 3
    assert spec.merges == [(0, 0, 1, 0), (0, 1, 0, 2)]
    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    table = document.tables[0]
    assert len(table.rows) == 3
    assert len(table.rows[0].cells) == 3
    assert table.cell(0, 0)._tc is table.cell(1, 0)._tc
    assert table.cell(0, 1)._tc is table.cell(0, 2)._tc


def test_docx_wide_table_sections_do_not_add_trailing_or_duplicate_transitions():
    def wide_table(value: str) -> str:
        header = "".join(f"<th>{value}{index}</th>" for index in range(10))
        row = "".join(f"<td>{index}</td>" for index in range(10))
        return f"<table><tr>{header}</tr><tr>{row}</tr></table>"

    wide = wide_table("wide")
    only_wide = Document(io.BytesIO(build_docx(sample_task(markdown=wide))))
    assert [str(section.orientation) for section in only_wide.sections] == [
        "PORTRAIT (0)",
        "LANDSCAPE (1)",
    ]

    adjacent_wide = Document(
        io.BytesIO(build_docx(sample_task(markdown=f"{wide}\n\n{wide}")))
    )
    assert [str(section.orientation) for section in adjacent_wide.sections] == [
        "PORTRAIT (0)",
        "LANDSCAPE (1)",
    ]


def test_latex_delimiters_are_removed_without_leaking_backslashes():
    text = plain_markdown_text(r"$(r,a)=\\arrayll R(q,r,a) \& iteration=0 \\ R(q,d_k^1,r,a)$")
    assert "$" not in text
    assert "\\" not in text
    assert "R(q,r,a)" in text


def test_docx_formulas_are_rasterized_with_array_layout_when_renderer_is_available():
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is None:
        pytest.skip("optional formula rasterizer is not installed in this test environment")
    markdown = r'''Inline $T_{l_k}^{\prime}=H_{l_k}\in R^{u\times d_{model}}$.

$$\mathrm{logit}_{l_k}=\mathrm{Prober}_{l_k}(T\prime_{l_k}) \tag{3}$$

$$ (r,\hat{a})=\left\{\begin{array}{ll}LLM.generate(q)&\text{if iteration}=0\\LLM.generate(q,\{d_k\}_{1}^{j})&\text{otherwise}\end{array}\right. \tag{5}$$

<table><tr><th>Formula</th><th>Value</th></tr><tr><td>$\theta$</td><td>$x^2+\frac{1}{N}$</td></tr></table>'''
    payload = build_docx(sample_task(markdown=markdown))
    document = Document(io.BytesIO(payload))
    # Three prose equations plus the two formula cells should be pictures;
    # no malformed literal array command or invisible OMML object is emitted.
    assert len(document.inline_shapes) >= 5
    assert document.tables[0].cell(1, 0).text == ""
    assert document.tables[0].cell(1, 1).text == ""
    with io.BytesIO(payload) as stream:
        import zipfile

        with zipfile.ZipFile(stream) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
    assert "arrayll" not in xml
    assert "oMath" not in xml
    assert "LaTeX formula:" in xml


def test_formula_image_renderer_handles_all_real_task_formulas():
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is None:
        pytest.skip("optional formula rasterizer is not installed in this test environment")
    import json
    import zipfile

    task_path = Path("data/tasks/mtmcje66-jzihwb/result.json")
    if not task_path.exists():
        pytest.skip("local real-task fixture is not present")
    task = json.loads(task_path.read_text(encoding="utf-8"))
    formulas = []
    for match in __import__("exporters").FORMULA_RE.finditer(task.get("markdown", "")):
        source = next(group for group in match.groups() if group is not None)
        if source not in formulas:
            formulas.append(source)
    assert formulas
    for source in formulas:
        image = _formula_image(source, display=True)
        assert image and image.startswith(b"\x89PNG\r\n\x1a\n")


def test_formula_images_keep_transparent_inset_for_table_cell_borders():
    math_to_image, _ = _load_formula_rasterizer()
    if math_to_image is None:
        pytest.skip("optional formula rasterizer is not installed in this test environment")
    payload = _formula_image(
        r"\sum_{i=1}^{N} y_i",
        display=False,
        table_cell=True,
    )
    assert payload
    image = Image.open(io.BytesIO(payload)).convert("RGBA")
    bbox = image.getchannel("A").getbbox()
    assert bbox is not None
    assert bbox[1] > 0
    assert bbox[3] < image.height


def test_docx_formulas_are_native_omml_and_handle_ocr_escaping():
    # The source mirrors the malformed doubled-command/array spelling emitted
    # by older OCR responses and the valid LaTeX returned by current adapters.
    malformed = r"\\left\\{\\begin{array}{ll}a&\\text{if }x=0\\\\b&\\text{otherwise}\\end{array}\\right."
    assert "\\\\text" in malformed
    normalized = _normalize_latex_source(malformed)
    assert r"\begin{array}{ll}" in normalized
    assert r"\text{if }" in normalized
    element = _formula_omml(normalized)
    assert element is not None
    xml = etree.tostring(element, encoding="unicode")
    assert "oMath" in xml
    assert "m:" in xml
    assert "arrayll" not in xml
    assert "text" not in xml


def test_docx_display_formula_is_rasterized_and_algorithm_lines_are_math():
    markdown = r"""# Formula smoke

$$L = -\frac{1}{N} \sum_{i=1}^{N} [y_i \log(p_i)]$$

Algorithm 1 Probing
1: Compute $(r,\hat{a})=\left\{\begin{array}{ll}q&\text{if }x=0\\r&\text{otherwise}\end{array}\right.$
Output: $\hat{a}$
    """
    document = Document(io.BytesIO(build_docx(sample_task(markdown=markdown))))
    math_to_image, _ = _load_formula_rasterizer()
    xml = document._element.xml
    # With the production rasterizer, display and algorithm equations are
    # pictures so Office suites do not silently drop OMML on import.
    if math_to_image is not None:
        assert len(document.inline_shapes) >= 3
        assert "oMath" not in xml
    # The algorithm's formula must never be emitted as literal OCR commands.
    assert "arrayll" not in xml


def test_docx_native_formula_mode_writes_editable_omml(monkeypatch):
    monkeypatch.setenv("PANDOCR_DOCX_FORMULA_MODE", "native")
    document = Document(io.BytesIO(build_docx(sample_task(markdown="$$x^2 + y^2$$"))))
    xml = document._element.xml
    assert "oMath" in xml
    assert "LaTeX formula:" not in xml
