"""Offline task exporters for editable Office files and searchable PDFs."""

from __future__ import annotations

import base64
import html
import io
import os
import re
import tempfile
import threading
from dataclasses import dataclass
from html.parser import HTMLParser
from functools import lru_cache
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageOps
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

try:
    from latex2mathml.converter import convert as _latex_to_mathml
    from mathml2omml import convert as _mathml_to_omml
except ImportError:  # pragma: no cover - exercised only in minimal installs
    _latex_to_mathml = None
    _mathml_to_omml = None


class ExportError(ValueError):
    """Raised when a task cannot be exported in the requested format."""


MARKDOWN_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+['\"][^'\"]*['\"])?\)")
HTML_IMAGE_LINE_RE = re.compile(
    r"^\s*(?:<div\b[^>]*>\s*)?<img\b(?P<attrs>[^>]*)/?>\s*(?:</div>)?\s*$",
    re.IGNORECASE | re.DOTALL,
)
HTML_IMAGE_ATTR_RE = re.compile(
    r"(?P<name>[A-Za-z_:][-A-Za-z0-9_:.]*)\s*=\s*"
    r"(?:\"(?P<double>[^\"]*)\"|'(?P<single>[^']*)'|(?P<bare>[^\s>]+))",
    re.IGNORECASE,
)
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_TOKEN_RE = re.compile(
    r"(`[^`\n]+`|\*\*[^*\n]+\*\*|(?<![\w_])__[^_\n]+__(?![\w_])|"
    r"(?<![\w*])\*[^*\n]+\*(?![\w*])|(?<![\w_])_[^_\n]+_(?![\w_]))"
)
HTML_TAG_RE = re.compile(
    r"</?(?:abbr|b|blockquote|br|code|del|details|div|em|h[1-6]|hr|i|kbd|li|mark|ol|p|pre|s|span|strong|sub|summary|sup|table|tbody|td|th|thead|tr|u|ul)(?:\s+[^<>]*?)?\s*/?>",
    flags=re.IGNORECASE,
)
FENCE_RE = re.compile(r"^\s*(```|~~~)")
HTML_TABLE_RE = re.compile(r"<table\b[^>]*>.*?</table>", re.IGNORECASE | re.DOTALL)
FORMULA_PREFIXES = ("=", "+", "-", "@")
LATEX_SYMBOLS = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "theta": "θ", "lambda": "λ", "mu": "μ", "pi": "π", "sigma": "σ",
    "phi": "φ", "omega": "ω", "times": "×", "cdot": "·", "pm": "±",
    "leq": "≤", "geq": "≥", "neq": "≠", "approx": "≈", "infty": "∞",
    "rightarrow": "→", "leftarrow": "←", "sum": "∑", "prod": "∏",
}
LATEX_COMMAND_NAMES = frozenset(
    {
        *LATEX_SYMBOLS,
        "text",
        "textbf",
        "textit",
        "textrm",
        "mathrm",
        "mathbf",
        "mathit",
        "mathsf",
        "mathtt",
        "mathbb",
        "mathcal",
        "mathfrak",
        "operatorname",
        "operatorname*",
        "frac",
        "dfrac",
        "tfrac",
        "sqrt",
        "root",
        "left",
        "right",
        "begin",
        "end",
        "tag",
        "displaystyle",
        "textstyle",
        "scriptstyle",
        "scriptstyle",
        "ldots",
        "cdots",
        "dots",
        "prime",
        "hat",
        "bar",
        "overline",
        "underline",
        "vec",
        "dot",
        "ddot",
        "boldsymbol",
        "arrayll",
    }
)
FORMULA_RE = re.compile(
    r"\$\$(?P<display>.+?)\$\$|\$(?P<inline>[^$\n]+?)\$|"
    r"\\\((?P<paren>.+?)\\\)|\\\[(?P<bracket>.+?)\\\]",
    re.DOTALL,
)
# A number of OCR/Markdown adapters emit display equations as a bare TeX
# environment instead of wrapping them in ``$``/``\[`` delimiters.  Keep this
# pattern separate from ``FORMULA_RE`` so the latter remains backwards
# compatible for callers that inspect its named groups.  The leading
# ``\\+`` accepts both normal TeX (``\\begin``) and JSON-doubled commands
# (``\\\\begin``); the source is normalized before rendering.
FORMULA_ENV_RE = re.compile(
    r"\\+begin\s*\{(?P<environment_name>"
    r"equation\*?|align\*?|aligned|gather\*?|multline\*?|"
    r"eqnarray\*?|displaymath\*?|split|array|cases|matrix|pmatrix|bmatrix"
    r")\}(?P<environment>.*?)\\+end\s*\{(?P=environment_name)\}",
    re.IGNORECASE | re.DOTALL,
)
# Older PaddleOCR-VL responses sometimes omit all math delimiters and use the
# compact ``\\arrayll ... \\array`` spelling.  It is intentionally line-scoped
# and requires the distinctive ``arrayll`` token; ordinary currency/prose is
# therefore not treated as a formula by this fallback detector.
FORMULA_SHORTHAND_RE = re.compile(
    r"(?P<shorthand>^[^\n\r]*?\\+arrayll\b[^\n\r]*$)",
    re.IGNORECASE | re.MULTILINE,
)
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
ARRAY_ENV_RE = re.compile(
    r"\\begin\s*\{(?P<env>array|cases|aligned|matrix|pmatrix|bmatrix)\}"
    r"(?:\s*\{[^{}]*\})?(?P<body>.*?)"
    r"\\end\s*\{(?P=env)\}",
    re.DOTALL,
)

# Formula images are used as the visual representation in DOCX.  OMML is
# still retained as a last-resort fallback for installations that do not have
# the rasterizer, but a raster image is what makes the export render
# consistently in WPS, Word, and LibreOffice.
FORMULA_IMAGE_DPI = 220
FORMULA_IMAGE_MAX_WIDTH_IN = 6.25
_MATH_RENDER_LOCK = threading.RLock()
_MATH_RENDERER_READY = False
_MATH_TO_IMAGE = None
_MATH_FONT_PROPERTIES = None


def _normalize_latex_source(source: object) -> str:
    """Normalize LaTeX variants commonly emitted by OCR/VLM adapters.

    OCR output is not always a TeX source verbatim.  In particular, JSON and
    Markdown serializers often double command backslashes (``\\\\alpha``),
    while some models emit the old ``\\arrayll ... \\array`` shorthand instead
    of a standard ``array`` environment.  Normalizing only command prefixes
    keeps genuine TeX row separators (``\\\\`` followed by whitespace) intact.
    """

    text = str(source or "").replace("\r\n", "\n").replace("\r", "\n")
    # A formula is a single math object in the DOCX.  Newlines inside it are
    # layout hints from OCR, not paragraph breaks.
    text = re.sub(r"\s*\n\s*", " ", text).strip()
    # Collapse doubled backslashes immediately before a command name.  Do this
    # repeatedly because a few adapters escape an already escaped command.
    previous = None
    while previous != text:
        previous = text
        # Collapse only known command names.  A broad ``\\\\(?=[a-z])`` rule
        # would incorrectly turn a row separator followed directly by a lower-
        # case matrix value (``a\\\\b``) into a command.
        command_pattern = "|".join(
            sorted((re.escape(name) for name in LATEX_COMMAND_NAMES), key=len, reverse=True)
        )
        text = re.sub(rf"\\\\(?=(?:{command_pattern})(?![A-Za-z]))", r"\\", text)
        # Escaped punctuation commands (``\\{``, ``\\[`` ...) are also
        # commonly doubled, whereas a bare ``\\\\`` remains an array row
        # separator.
        text = re.sub(r"\\\\(?=[{}\[\]()|])", r"\\", text)

    # OCR occasionally escapes an ordinary one-letter variable as if it were
    # a TeX command (for example ``\\d_k`` or ``\\x^2``).  Matplotlib then
    # interprets the letter as an accent command and the subscript/superscript
    # is rendered as literal text.  In a math source, a one-letter token
    # immediately followed by ``_``/``^`` is overwhelmingly a variable; strip
    # the spurious escape.  Apply the doubled form first and repeat so JSON-
    # escaped input (``\\\\d_k``) is handled as well.  Named commands such as
    # ``\\hat{a}``/``\\text{...}`` are unaffected because they are not followed
    # directly by a subscript or superscript marker.
    previous = None
    while previous != text:
        previous = text
        text = re.sub(r"\\\\(?P<letter>[A-Za-z])(?=[_^])", r"\g<letter>", text)
        text = re.sub(r"(?<!\\)\\(?P<letter>[A-Za-z])(?=[_^])", r"\g<letter>", text)
    # Markdown/OCR escaping can also leave a slash in front of a math
    # operator.  Keep the operator while preserving genuine row separators.
    text = re.sub(r"(?<!\\)\\([_^])", r"\1", text)

    # PaddleOCR-VL may insert spaces between letters inside an upright
    # operator, e.g. ``\\mathrm{m o d e l}`` or
    # ``\\operatorname{N o r m}``.  Those spaces are OCR artifacts, not
    # intended typography; collapsing them keeps identifiers readable while
    # leaving ``\\text{if iteration = 0}`` prose untouched.
    spaced_identifier_commands = r"mathrm|operatorname\*?|mathbf|mathit|mathsf|mathtt|mathbb|mathcal|mathfrak"
    text = re.sub(
        rf"(\\(?:{spaced_identifier_commands})\s*\{{)([A-Za-z](?:\s+[A-Za-z])+)(\}})",
        lambda match: match.group(1) + re.sub(r"\s+", "", match.group(2)) + match.group(3),
        text,
    )

    # Normalize malformed array shorthand seen in older PaddleOCR-VL output.
    # ``\\arrayll`` is equivalent to ``\\begin{array}{ll}``; a trailing
    # ``\\array`` is its corresponding end marker.
    text = re.sub(r"\\arrayll\b", r"\\begin{array}{ll}", text)
    if "\\begin{array}{ll}" in text and "\\end{array}" not in text:
        text = re.sub(r"\\array\s*$", r"\\end{array}", text)

    # A few models emit the environment without the braces around its column
    # specification.  The converter accepts the canonical form consistently.
    text = re.sub(r"\\begin\{array\}\s*ll\b", r"\\begin{array}{ll}", text)
    return text


def _backslash_run_before(text: str, index: int) -> int:
    """Return the number of consecutive backslashes immediately before *index*."""

    count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        count += 1
        cursor -= 1
    return count


def _is_escaped_delimiter(text: str, index: int) -> bool:
    """Whether the delimiter beginning at *index* is escaped by an odd slash run."""

    return _backslash_run_before(text, index) % 2 == 1


def _formula_source_from_match(match: re.Match) -> str:
    """Extract a renderable formula source from a standard or environment match."""

    groups = match.groupdict()
    shorthand = groups.get("shorthand")
    if shorthand is not None:
        return _shorthand_formula_source(str(shorthand))
    environment_body = groups.get("environment")
    if environment_body is not None:
        environment_name = str(groups.get("environment_name") or "").lower().rstrip("*")
        body = str(environment_body)
        tag_suffix = ""
        tag_match = re.search(r"\\tag\s*\{([^{}]+)\}\s*$", body)
        if tag_match:
            tag_suffix = rf" \tag{{{tag_match.group(1).strip()}}}"
            body = body[: tag_match.start()].rstrip()
        # ``equation`` is already a single math object.  Multi-line TeX
        # environments need a lightweight array wrapper so the same renderer
        # can align rows and columns without leaking literal ``&``/``\\``.
        if environment_name in {"align", "aligned", "gather", "multline", "eqnarray", "displaymath", "split"}:
            return rf"\begin{{array}}{{l}}{body}\end{{array}}" + tag_suffix
        if environment_name == "array":
            # The column specification follows ``\\begin{array}`` and is
            # therefore part of the captured body.  Remove it before adding a
            # canonical spec understood by the manual array compositor.
            body = re.sub(r"^\s*\{[^{}]*\}", "", body, count=1)
            return rf"\begin{{array}}{{ll}}{body}\end{{array}}" + tag_suffix
        if environment_name in {"cases", "matrix", "pmatrix", "bmatrix"}:
            return rf"\begin{{{environment_name}}}{body}\end{{{environment_name}}}" + tag_suffix
        return body + tag_suffix
    for name in ("display", "inline", "paren", "bracket"):
        value = groups.get(name)
        if value is not None:
            return str(value)
    return ""


def _shorthand_formula_source(value: str) -> str:
    """Convert an OCR ``arrayll`` line into a canonical display formula."""

    raw = str(value or "").strip()
    command = re.search(r"\\+arrayll\b", raw, flags=re.IGNORECASE)
    if command is None:
        return ""
    prefix = raw[: command.start()].strip()
    # A line may be introduced as ``Equation:``/``Formula:``.  Keep only the
    # mathematical suffix before the shorthand marker; colons are not valid
    # TeX in this compact form and otherwise become visible in the image.
    if re.search(r"[:：]", prefix):
        prefix = re.split(r"[:：]", prefix)[-1].strip()
    # Do not consume arbitrary prose merely because it contains the legacy
    # token.  A valid prefix is a compact mathematical expression such as
    # ``(r,a)=`` or ``\\left\\{``; a sentence like ``No delimiter`` is not.
    if prefix and re.search(r"\b[A-Za-z]{3,}\b\s+\b[A-Za-z]{2,}\b", prefix):
        return ""
    if prefix and not re.search(r"[=^_{}()[\]+\-*/,.]|\\", prefix):
        return ""
    remainder = raw[command.end() :].strip()
    if "$" in prefix or "$" in remainder:
        # A properly delimited formula is handled by FORMULA_RE.  Refuse this
        # fallback candidate rather than consuming surrounding prose.
        return ""

    tag = ""
    label = re.search(r"(?:^|\s)\(\s*(\d+)\s*\)\s*[.,;:]?\s*$", remainder)
    if label:
        tag = label.group(1)
        remainder = remainder[: label.start()].rstrip()
    # OCR may leave punctuation after the array terminator or omit the
    # terminator entirely.  In both cases the canonical wrapper below closes
    # the environment deterministically.
    remainder = re.sub(r"\s*\\+\s*right\s*\.?\s*$", "", remainder, flags=re.IGNORECASE)
    remainder = re.sub(r"\s*\\+\s*array\s*$", "", remainder, flags=re.IGNORECASE)
    remainder = re.sub(r"\s*\\+\s*end\s*\{array\}\s*$", "", remainder, flags=re.IGNORECASE)
    remainder = remainder.rstrip(" \t.,;:")
    source = f"{prefix} \\begin{{array}}{{ll}} {remainder} \\end{{array}}"
    if tag:
        source += rf" \tag{{{tag}}}"
    return _normalize_latex_source(source)


def _formula_match_is_display(match: re.Match) -> bool:
    groups = match.groupdict()
    return bool(
        groups.get("shorthand") is not None
        or groups.get("environment") is not None
        or groups.get("display") is not None
        or groups.get("bracket") is not None
    )


def _looks_like_formula(source: str, *, display: bool = False) -> bool:
    """Reject accidental currency/ prose dollar pairs while keeping math useful."""

    raw_source = str(source or "")
    raw_value = raw_source.strip()
    value = _normalize_latex_source(raw_value).strip()
    if not value:
        return False
    # Display delimiters and explicit TeX environments are an unambiguous
    # author/OCR signal, even when the body is a short single variable.
    if display:
        return True
    if re.search(r"\\[A-Za-z]+", value):
        return True
    if re.search(r"[_^=+\-*/<>≤≥≠≈×÷±∑∏∫√∈→←]", value):
        return True
    # Short variable/function/number forms such as ``$x$``, ``$O(n)$`` and
    # ``$10$`` are valid inline math.  Ignore whitespace at the delimiters
    # (OCR commonly emits ``$ (r) $``), but reject internal whitespace so a
    # prose/currency span like ``Cost is $5 and tax is $10`` is not consumed.
    if len(value) <= 64 and not re.search(r"\s", value):
        if re.fullmatch(r"[A-Za-z0-9_.,'()\[\]{}]+", value):
            return True
    return False


def _iter_formula_matches(text: object):
    """Yield valid, non-overlapping formula matches in source order.

    ``re.finditer`` alone cannot distinguish an escaped ``$`` from a math
    delimiter and may consume a later valid pair while rejecting an earlier
    escaped pair.  This small scanner retries from the next character when a
    candidate is rejected, preserving both literal dollars and following
    formulas.
    """

    value = str(text or "")
    candidates: list[re.Match] = []
    position = 0
    while position < len(value):
        match = FORMULA_RE.search(value, position)
        if match is None:
            break
        groups = match.groupdict()
        source = _formula_source_from_match(match)
        valid = True
        if groups.get("display") is not None or groups.get("inline") is not None:
            delimiter_length = 2 if groups.get("display") is not None else 1
            closing_start = match.end() - delimiter_length
            if _is_escaped_delimiter(value, match.start()) or _is_escaped_delimiter(value, closing_start):
                valid = False
            elif not _looks_like_formula(source, display=groups.get("display") is not None):
                valid = False
        elif groups.get("paren") is not None or groups.get("bracket") is not None:
            opening_start = match.start()
            closing_start = match.end() - 2
            if _is_escaped_delimiter(value, opening_start) or _is_escaped_delimiter(value, closing_start):
                valid = False
        if valid:
            candidates.append(match)
            position = match.end()
        else:
            position = match.start() + 1

    for match in FORMULA_ENV_RE.finditer(value):
        if _is_escaped_delimiter(value, match.start()):
            continue
        source = _formula_source_from_match(match)
        if _looks_like_formula(source, display=True):
            candidates.append(match)

    for match in FORMULA_SHORTHAND_RE.finditer(value):
        raw = str(match.group("shorthand") or "")
        # Require either an explicit terminator, an alignment marker, or a row
        # separator.  This keeps a stray prose/code token named ``arrayll``
        # from turning an entire line into an image.
        if not re.search(r"\\+\s*(?:array|right)\b|&|\\{2,}|[_^=+\-*/<>]", raw, flags=re.IGNORECASE):
            continue
        source = _formula_source_from_match(match)
        if source and _looks_like_formula(source, display=True):
            candidates.append(match)

    last_end = -1
    for match in sorted(candidates, key=lambda item: (item.start(), -item.end())):
        if match.start() < last_end:
            continue
        last_end = match.end()
        yield match


def _preferred_cjk_font() -> str:
    configured = os.getenv("PANDOCR_DOCX_CJK_FONT", "").strip()
    if configured:
        return configured
    candidates = (
        (Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"), "Noto Sans CJK SC"),
        (Path("/System/Library/Fonts/Hiragino Sans GB.ttc"), "Hiragino Sans GB"),
        (Path("/Library/Fonts/Arial Unicode.ttf"), "Arial Unicode MS"),
        (Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"), "Arial Unicode MS"),
        (Path("C:/Windows/Fonts/msyh.ttc"), "Microsoft YaHei"),
    )
    return next((font_name for path, font_name in candidates if path.exists()), "Arial Unicode MS")


CJK_FONT = _preferred_cjk_font()


def normalize_markdown(value: object) -> str:
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def _pipe_is_escaped(text: str, index: int) -> bool:
    backslashes = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        backslashes += 1
        cursor -= 1
    return backslashes % 2 == 1


def split_markdown_table_row(line: str) -> list[str]:
    text = str(line or "").strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|") and not _pipe_is_escaped(text, len(text) - 1):
        text = text[:-1]
    cells: list[str] = []
    cell: list[str] = []
    for character in text:
        if character == "|":
            trailing_backslashes = 0
            for current in reversed(cell):
                if current != "\\":
                    break
                trailing_backslashes += 1
            if trailing_backslashes % 2 == 1:
                cell.pop()
                cell.append("|")
            else:
                cells.append("".join(cell).strip())
                cell = []
        else:
            cell.append(character)
    cells.append("".join(cell).strip())
    return cells


def is_markdown_table_separator(line: str) -> bool:
    cells = split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", re.sub(r"\s+", "", cell)) for cell in cells)


def extract_markdown_tables(markdown: str) -> list[list[list[str]]]:
    lines = normalize_markdown(markdown).split("\n")
    tables: list[list[list[str]]] = []
    in_fence = False
    index = 0
    while index < len(lines) - 1:
        if FENCE_RE.match(lines[index]):
            in_fence = not in_fence
            index += 1
            continue
        if in_fence or "|" not in lines[index] or not is_markdown_table_separator(lines[index + 1]):
            index += 1
            continue
        header = split_markdown_table_row(lines[index])
        if len(header) < 2:
            index += 1
            continue
        rows = [header]
        index += 2
        while index < len(lines) and "|" in lines[index] and not FENCE_RE.match(lines[index]):
            row = split_markdown_table_row(lines[index])
            if len(row) < 2:
                break
            rows.append((row + [""] * len(header))[: len(header)])
            index += 1
        tables.append(rows)
    return tables


@dataclass(frozen=True)
class HtmlTableSpec:
    rows: list[list[str]]
    merges: list[tuple[int, int, int, int]]
    header_rows: int = 1


def _layout_html_table(
    raw_rows: list[list[tuple[str, int, int]]],
) -> HtmlTableSpec | None:
    """Expand HTML row/column spans into a rectangular Office table grid."""

    if not raw_rows:
        return None
    occupied: set[tuple[int, int]] = set()
    values: dict[tuple[int, int], str] = {}
    merges: list[tuple[int, int, int, int]] = []
    row_has_span: list[bool] = []
    max_row = len(raw_rows) - 1
    max_column = -1
    for row_index, cells in enumerate(raw_rows):
        column_index = 0
        row_has_span.append(False)
        for value, rowspan, colspan in cells:
            while (row_index, column_index) in occupied:
                column_index += 1
            rowspan = max(1, int(rowspan or 1))
            colspan = max(1, int(colspan or 1))
            values[(row_index, column_index)] = value
            for target_row in range(row_index, row_index + rowspan):
                for target_column in range(column_index, column_index + colspan):
                    occupied.add((target_row, target_column))
            if rowspan > 1 or colspan > 1:
                merges.append(
                    (
                        row_index,
                        column_index,
                        row_index + rowspan - 1,
                        column_index + colspan - 1,
                    )
                )
                row_has_span[-1] = True
            max_row = max(max_row, row_index + rowspan - 1)
            max_column = max(max_column, column_index + colspan - 1)
            column_index += colspan
    if max_column < 0:
        return None
    rows = [
        [values.get((row, column), "") for column in range(max_column + 1)]
        for row in range(max_row + 1)
    ]
    header_rows = 0
    for has_span in row_has_span:
        if not has_span:
            break
        header_rows += 1
    header_rows = min(max(1, header_rows), len(rows))
    return HtmlTableSpec(rows=rows, merges=merges, header_rows=header_rows)


class _HtmlTableParser(HTMLParser):
    """Extract OCR-produced HTML tables without requiring BeautifulSoup."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tables: list[HtmlTableSpec] = []
        self.rows: list[list[tuple[str, int, int]]] = []
        self.current_row: list[tuple[str, int, int]] | None = None
        self.current_cell: list[str] | None = None
        self.current_rowspan = 1
        self.current_colspan = 1

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag == "tr":
            self.current_row = []
        elif tag == "br" and self.current_cell is not None:
            # ``HTMLParser`` does not emit a data callback for ``<br>``.  Add
            # the separator explicitly or adjacent cell lines collapse into
            # one word before the DOCX/XLSX writers see them.
            self.current_cell.append("\n")
        elif tag in {"td", "th"} and self.current_row is not None:
            self.current_cell = []
            attributes = {str(name).lower(): value for name, value in attrs}
            try:
                self.current_rowspan = max(1, int(attributes.get("rowspan") or 1))
            except (TypeError, ValueError):
                self.current_rowspan = 1
            try:
                self.current_colspan = max(1, int(attributes.get("colspan") or 1))
            except (TypeError, ValueError):
                self.current_colspan = 1

    def handle_data(self, data: str) -> None:
        if self.current_cell is not None:
            self.current_cell.append(data)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"td", "th"} and self.current_row is not None and self.current_cell is not None:
            self.current_row.append(
                (
                    _plain_markdown_preserving_formulas("".join(self.current_cell)),
                    self.current_rowspan,
                    self.current_colspan,
                )
            )
            self.current_cell = None
            self.current_rowspan = 1
            self.current_colspan = 1
        elif tag == "tr" and self.current_row is not None:
            if self.current_row:
                self.rows.append(self.current_row)
            self.current_row = None
        elif tag == "table":
            table = _layout_html_table(self.rows)
            if table is not None:
                self.tables.append(table)
            self.rows = []


def extract_html_table_specs(markdown: str) -> list[HtmlTableSpec]:
    parser = _HtmlTableParser()
    for match in HTML_TABLE_RE.finditer(normalize_markdown(markdown)):
        parser.feed(match.group(0))
    return parser.tables


def extract_html_tables(markdown: str) -> list[list[list[str]]]:
    return [table.rows for table in extract_html_table_specs(markdown)]


def _prepare_html_table_markers(markdown: str) -> tuple[str, list[HtmlTableSpec]]:
    tables: list[HtmlTableSpec] = []

    def replace(match: re.Match) -> str:
        parsed = extract_html_table_specs(match.group(0))
        if not parsed:
            return match.group(0)
        index = len(tables)
        tables.extend(parsed)
        return f"\n__PANDOCR_HTML_TABLE_{index}__\n"

    return HTML_TABLE_RE.sub(replace, normalize_markdown(markdown)), tables


def _inline_token_content(match: re.Match) -> str:
    token = match.group(0)
    delimiter_length = 2 if token.startswith(("**", "__")) else 1
    return token[delimiter_length:-delimiter_length]


def _plain_markdown_preserving_formulas(value: object) -> str:
    """Strip HTML/Markdown decoration without destroying math delimiters.

    HTML tables are parsed before DOCX/XLSX generation.  Calling
    :func:`plain_markdown_text` directly at that stage used to remove ``$``
    delimiters, making a formula cell impossible to convert to OMML later.
    """

    raw = str(value or "")
    pieces: list[str] = []
    position = 0
    for match in _iter_formula_matches(raw):
        if match.start() > position:
            pieces.append(plain_markdown_text(raw[position : match.start()]))
        pieces.append(match.group(0))
        position = match.end()
    if position < len(raw):
        pieces.append(plain_markdown_text(raw[position:]))
    return "".join(pieces).strip()


def plain_markdown_text(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = MARKDOWN_IMAGE_RE.sub(lambda match: match.group(1), text)
    text = MARKDOWN_LINK_RE.sub(lambda match: match.group(1), text)
    text = INLINE_TOKEN_RE.sub(_inline_token_content, text)
    text = HTML_TAG_RE.sub("", text)
    text = html.unescape(text)
    # Replace only validated formulas.  Earlier code treated *any* dollar sign
    # as a math trigger, which silently removed currency symbols and could turn
    # an entire sentence between two dollars into a formula.  Escaped dollars
    # are left in the surrounding text and are unescaped below.
    formula_pieces: list[str] = []
    formula_position = 0
    for match in _iter_formula_matches(text):
        if match.start() > formula_position:
            formula_pieces.append(text[formula_position : match.start()])
        formula_pieces.append(_formula_fallback_text(_formula_source_from_match(match)))
        formula_position = match.end()
    if formula_pieces:
        formula_pieces.append(text[formula_position:])
        text = "".join(formula_pieces)

    # Markdown uses a backslash to escape a literal pipe or dollar.  Keep other
    # backslashes intact because OCR commonly contains Windows paths and LaTeX
    # commands (for example ``C:\\temp`` and ``\\alpha``).  Counting
    # consecutive slashes also preserves the even/odd table-delimiter rule.
    unescaped: list[str] = []
    backslashes = 0
    for character in text:
        if character == "\\":
            backslashes += 1
            continue
        if character in {"|", "$"} and backslashes % 2:
            unescaped.extend("\\" * (backslashes - 1))
        else:
            unescaped.extend("\\" * backslashes)
        unescaped.append(character)
        backslashes = 0
    unescaped.extend("\\" * backslashes)
    text = "".join(unescaped)
    return text.strip()


def _formula_parts(source: str) -> tuple[str, str]:
    """Return LaTeX body and an optional equation number from a formula."""
    value = _normalize_latex_source(source)
    tag = ""
    tag_match = re.search(r"\\tag\*?\s*\{([^{}]+)\}\s*$", value)
    if tag_match:
        tag = tag_match.group(1).strip()
        if len(tag) >= 2 and tag.startswith("(") and tag.endswith(")"):
            tag = tag[1:-1].strip()
        value = value[: tag_match.start()].rstrip()
    value = re.sub(r"^\\(?:displaystyle|textstyle|scriptstyle)\s+", "", value)
    return value, tag


def _formula_fallback_text(source: str) -> str:
    """Readable fallback when an Office math converter is unavailable."""
    body, tag = _formula_parts(source)
    # Row separators are layout markers, not literal backslashes.  Remove them
    # before command substitution so a row beginning with ``b`` cannot leave a
    # stray ``\\b`` in the fallback text.
    text = re.sub(r"\\\\+", " ", body)
    # Strip environment wrappers as a unit; removing only ``\\begin`` leaves
    # confusing remnants such as ``beginarrayll`` in searchable text.
    text = re.sub(r"\\begin\s*\{[^{}]*\}(?:\s*\{[^{}]*\})?", "", text)
    text = re.sub(r"\\end\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\(?:text|mathrm|mathbf|mathit|operatorname)\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\([A-Za-z]+)", lambda match: LATEX_SYMBOLS.get(match.group(1), match.group(1)), text)
    text = re.sub(r"\\(?:left|right)\b", "", text)
    text = text.replace("\\{", "{").replace("\\}", "}")
    text = text.replace("\\,", " ").replace("\\;", " ").replace("\\!", "")
    text = re.sub(r"\{([^{}]*)\}", r"\1", text)
    text = text.replace("&", " ").replace("$", "")
    # Any remaining slash is an unsupported TeX escape.  In a fallback string
    # it is more useful (and less confusing) to show the symbol than a command
    # escape, and this function is only called for formula sources.
    text = text.replace("\\", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return f"{text} ({tag})" if tag else text


def _load_formula_rasterizer():
    """Load Matplotlib's isolated mathtext renderer on first use.

    Exporting a document should not make the web process depend on a GUI
    backend.  ``Agg`` is explicitly selected and Matplotlib's cache is pointed
    at the container's writable temporary directory because the production
    image is read-only apart from ``/tmp``.
    """

    global _MATH_RENDERER_READY, _MATH_TO_IMAGE, _MATH_FONT_PROPERTIES
    if _MATH_RENDERER_READY:
        return _MATH_TO_IMAGE, _MATH_FONT_PROPERTIES
    with _MATH_RENDER_LOCK:
        if _MATH_RENDERER_READY:
            return _MATH_TO_IMAGE, _MATH_FONT_PROPERTIES
        _MATH_RENDERER_READY = True
        try:
            cache_dir = Path(tempfile.gettempdir()) / "pandocr-matplotlib"
            cache_dir.mkdir(parents=True, exist_ok=True)
            os.environ.setdefault("MPLCONFIGDIR", str(cache_dir))
            os.environ.setdefault("MPLBACKEND", "Agg")
            import matplotlib

            matplotlib.use("Agg", force=True)
            from matplotlib.font_manager import FontProperties
            from matplotlib.mathtext import math_to_image

            _MATH_TO_IMAGE = math_to_image
            _MATH_FONT_PROPERTIES = FontProperties
        except Exception:
            # The exporter remains usable in a deliberately minimal image;
            # callers then fall back to OMML/plain text below.
            _MATH_TO_IMAGE = None
            _MATH_FONT_PROPERTIES = None
    return _MATH_TO_IMAGE, _MATH_FONT_PROPERTIES


def _mathtext_source(source: str) -> str:
    """Turn common OCR LaTeX variants into Matplotlib mathtext syntax."""

    value = _normalize_latex_source(source).strip()
    # ``math_to_image`` supplies the outer dollar delimiters.  Strip accidental
    # delimiters that can occur when an adapter nests Markdown in a cell.
    value = re.sub(r"^\$+", "", value)
    value = re.sub(r"\$+$", "", value)
    value = re.sub(r"\\(?:displaystyle|textstyle|scriptstyle)\b", "", value)
    value = value.replace(r"\dfrac", r"\frac").replace(r"\tfrac", r"\frac")
    # Matplotlib supports ``\text`` and ``\operatorname`` in current
    # releases.  Mapping the latter to ``\mathrm`` keeps compatibility with
    # older mathtext versions shipped by some downstream images.
    value = re.sub(r"\\operatorname\*?\s*", lambda _match: r"\mathrm ", value)
    value = re.sub(r"\\dim\b", lambda _match: r"\mathrm{dim}", value)
    value = value.replace(r"\left.", "").replace(r"\right.", "")
    value = re.sub(r"\\left\s*([\[({])", r"\1", value)
    value = re.sub(r"\\right\s*([\])}])", r"\1", value)
    # A few OCR outputs contain a literal non-breaking space in ``\text``;
    # normalizing it avoids a parser error while retaining readable spacing.
    value = value.replace("\u00a0", " ")
    return value.strip()


def _split_formula_rows(value: str) -> list[str]:
    """Split an array body on TeX row separators without touching commands."""

    rows: list[str] = []
    current: list[str] = []
    index = 0
    while index < len(value):
        if value[index] == "\\":
            run_end = index
            while run_end < len(value) and value[run_end] == "\\":
                run_end += 1
            run_length = run_end - index
            # A slash run directly before ``&`` is an escaped alignment
            # marker emitted by some Markdown/OCR serializers, not a row
            # break.  Leave the marker for ``_split_formula_cells``.
            if run_length >= 1 and run_end < len(value) and value[run_end] == "&":
                current.append("&")
                index = run_end + 1
                continue
            if run_length >= 2:
                rows.append("".join(current).strip())
                current = []
                index = run_end
                while index < len(value) and value[index].isspace():
                    index += 1
                continue
        current.append(value[index])
        index += 1
    rows.append("".join(current).strip())
    return [row for row in rows if row]


def _split_formula_cells(value: str) -> list[str]:
    """Split an array row at unescaped alignment ampersands."""

    cells: list[str] = []
    current: list[str] = []
    index = 0
    while index < len(value):
        character = value[index]
        if character == "&":
            slash_count = 0
            cursor = index - 1
            while cursor >= 0 and value[cursor] == "\\":
                slash_count += 1
                cursor -= 1
            if slash_count % 2 == 0:
                cells.append("".join(current).strip())
                current = []
                index += 1
                continue
        current.append(character)
        index += 1
    cells.append("".join(current).strip())
    return cells


def _render_formula_text_image(text: str, *, font_size: int = 24) -> Image.Image:
    """Render a readable text fallback when mathtext rejects a formula."""

    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
    )
    font = None
    for candidate in candidates:
        if candidate.exists():
            try:
                from PIL import ImageFont

                font = ImageFont.truetype(str(candidate), font_size)
                break
            except Exception:
                continue
    if font is None:
        from PIL import ImageFont

        font = ImageFont.load_default()
    display = _formula_fallback_text(text)
    probe = Image.new("RGBA", (8, 8), (255, 255, 255, 0))
    draw = ImageDraw.Draw(probe)
    bbox = draw.textbbox((0, 0), display or " ", font=font)
    width = max(8, bbox[2] - bbox[0] + 4)
    height = max(8, bbox[3] - bbox[1] + 4)
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    ImageDraw.Draw(image).text((2 - bbox[0], 2 - bbox[1]), display or " ", fill=(11, 37, 69, 255), font=font)
    return image


def _render_mathtext_piece(source: str, *, font_size: float = 12.0) -> Image.Image:
    """Render one non-array math fragment to a transparent PNG image."""

    math_to_image, font_properties = _load_formula_rasterizer()
    if math_to_image is None or font_properties is None:
        return _render_formula_text_image(source)
    value = _mathtext_source(source)
    if not value:
        value = r"\;"
    try:
        stream = io.BytesIO()
        prop = font_properties(size=font_size)
        with _MATH_RENDER_LOCK:
            math_to_image(
                f"${value}$",
                stream,
                prop=prop,
                dpi=FORMULA_IMAGE_DPI,
                format="png",
                color="#0B2545",
            )
        stream.seek(0)
        image = Image.open(stream).convert("RGBA")
        # Detach from the BytesIO object before it goes out of scope.
        image.load()
        return image
    except Exception:
        return _render_formula_text_image(source)


def _compose_array_formula(source: str, *, font_size: float = 12.0) -> Image.Image | None:
    """Render an ``array``/``cases`` formula as a neatly aligned image."""

    match = ARRAY_ENV_RE.search(_normalize_latex_source(source))
    if not match:
        return None
    value = _normalize_latex_source(source)
    prefix = value[: match.start()]
    suffix = value[match.end() :]
    body = match.group("body")
    rows: list[Image.Image] = []
    for row in _split_formula_rows(body):
        row = row.replace(r"\hline", "").strip()
        if not row:
            continue
        cells = _split_formula_cells(row)
        cell_images = [_render_mathtext_piece(cell or r"\;", font_size=font_size) for cell in cells]
        row_height = max(image.height for image in cell_images)
        row_width = sum(image.width for image in cell_images) + max(0, len(cell_images) - 1) * 24
        row_image = Image.new("RGBA", (row_width, row_height), (255, 255, 255, 0))
        cursor = 0
        for image in cell_images:
            row_image.alpha_composite(image, (cursor, (row_height - image.height) // 2))
            cursor += image.width + 24
        rows.append(row_image)
    if not rows:
        return None

    row_gap = max(5, round(font_size * 0.45))
    rows_height = sum(image.height for image in rows) + row_gap * (len(rows) - 1)
    rows_width = max(image.width for image in rows)
    stack = Image.new("RGBA", (rows_width, rows_height), (255, 255, 255, 0))
    cursor_y = 0
    for image in rows:
        stack.alpha_composite(image, (0, cursor_y))
        cursor_y += image.height + row_gap

    # ``\left\{`` is not accepted as a standalone mathtext fragment.  A
    # regular escaped brace is, so stretch that glyph to the array height.
    has_left_brace = bool(re.search(r"\\left\s*\\\{", prefix)) or match.group("env") == "cases"
    parts: list[Image.Image] = []
    clean_prefix = re.sub(r"\\left\s*\\\{", "", prefix)
    clean_prefix = re.sub(r"\\left\s*([\[({])", r"\1", clean_prefix).strip()
    clean_suffix = re.sub(r"\\right\s*\.", "", suffix)
    clean_suffix = re.sub(r"\\right\s*([\])}])", r"\1", clean_suffix).strip()
    if clean_prefix:
        parts.append(_render_mathtext_piece(clean_prefix, font_size=font_size))
    if has_left_brace:
        brace = _render_mathtext_piece(r"\{", font_size=max(22, font_size * 2.5))
        brace_width = max(10, brace.width)
        brace = brace.resize((brace_width, rows_height), Image.Resampling.LANCZOS)
        parts.append(brace)
    parts.append(stack)
    if clean_suffix:
        parts.append(_render_mathtext_piece(clean_suffix, font_size=font_size))
    gap = max(10, round(font_size * 0.8))
    width = sum(image.width for image in parts) + gap * (len(parts) - 1)
    height = max(image.height for image in parts)
    output = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    cursor_x = 0
    for image in parts:
        output.alpha_composite(image, (cursor_x, (height - image.height) // 2))
        cursor_x += image.width + gap
    return output


@lru_cache(maxsize=512)
def _formula_image(
    source: str,
    display: bool = False,
    table_cell: bool = False,
) -> bytes | None:
    """Return a cached transparent PNG for a LaTeX formula."""

    body, _ = _formula_parts(source)
    if not body.strip():
        return None
    image = _compose_array_formula(body, font_size=14.0 if display else 11.5)
    if image is None:
        image = _render_mathtext_piece(body, font_size=14.0 if display else 11.5)
    if image is None:
        return None
    # ``math_to_image`` tightly crops the glyph bounding box.  Normal Word
    # paragraphs already provide enough line leading, while table cells use
    # single-line spacing and can otherwise put a subscript directly against
    # the lower border.  Add vertical breathing room only inside table cells;
    # keeping ordinary equations unchanged avoids moving equation numbers or
    # perturbing page breaks in long reports.
    if table_cell:
        image = ImageOps.expand(
            image,
            border=(0, 6, 0, 6),
            fill=(255, 255, 255, 0),
        )
    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def _set_picture_description(run, source: str) -> None:
    """Attach the original formula as accessible metadata on a picture."""

    drawing_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for element in run._r.iter():
        if element.tag == f"{{{drawing_ns}}}docPr":
            element.set("title", "Formula")
            element.set("descr", f"LaTeX formula: {_normalize_latex_source(source)}")


def _append_formula_image(
    paragraph,
    source: str,
    *,
    display: bool = False,
    table_cell: bool = False,
) -> bool:
    payload = _formula_image(source, display=display, table_cell=table_cell)
    if not payload:
        return False
    try:
        with Image.open(io.BytesIO(payload)) as image:
            width_px, height_px = image.size
        width_in = min(FORMULA_IMAGE_MAX_WIDTH_IN, max(0.08, width_px / FORMULA_IMAGE_DPI))
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(payload), width=Inches(width_in))
        _set_picture_description(run, source)
        return True
    except Exception:
        return False


def _formula_omml(source: str, *, display: bool = False):
    """Build a real Word math object from a LaTeX formula, if supported."""
    if not (_latex_to_mathml and _mathml_to_omml):
        return None
    body, _ = _formula_parts(source)
    try:
        omml = _mathml_to_omml(_latex_to_mathml(body))
        if not omml or "<m:oMath" not in omml:
            return None
        if "xmlns:m=" not in omml:
            omml = omml.replace("<m:oMath>", f'<m:oMath xmlns:m="{MATH_NS}">', 1)
        # Parse the fragment under a temporary root.  The converter emits an
        # ``m:oMath`` element with a namespace prefix but no declaration.
        # Parsing through a temporary root is more reliable than injecting an
        # xmlns attribute into the first tag (and also handles future converter
        # wrappers).
        if "xmlns:m=" not in omml:
            omml = f'<m:root xmlns:m="{MATH_NS}">{omml}</m:root>'
        root = parse_xml(omml)
        if root.tag == f"{{{MATH_NS}}}root":
            root = next((child for child in root if child.tag == f"{{{MATH_NS}}}oMath"), None)
            if root is None:
                return None
        # mathml2omml emits a document-level m:box wrapper. Word accepts it,
        # but LibreOffice/WPS may treat that wrapper as an empty object. Keep
        # the actual math children directly under m:oMath for interoperability.
        if len(root) == 1 and root[0].tag == f"{{{MATH_NS}}}box":
            box = root[0]
            if len(box) == 1 and box[0].tag == f"{{{MATH_NS}}}e":
                root[:] = list(box[0])
        if display:
            wrapper = OxmlElement("m:oMathPara")
            wrapper.append(root)
            return wrapper
        return root
    except (Exception,):
        return None


def _append_formula(
    paragraph,
    source: str,
    *,
    display: bool = False,
    table_cell: bool = False,
) -> bool:
    # Native OMML is the preferred DOCX representation when explicitly enabled
    # by the application.  The image path remains the compatibility fallback
    # for LibreOffice/WPS builds that cannot render third-party OMML reliably.
    native_mode = os.getenv("PANDOCR_DOCX_FORMULA_MODE", "image").strip().lower() in {
        "native", "omml", "editable"
    }
    if native_mode:
        element = _formula_omml(source, display=display)
        if element is not None:
            paragraph._p.append(element)
            _, tag = _formula_parts(source)
            if tag:
                run = paragraph.add_run(f" ({tag})")
                _set_run_font(run)
            return True
    if _append_formula_image(
        paragraph,
        source,
        display=display,
        table_cell=table_cell,
    ):
        _, tag = _formula_parts(source)
        if tag:
            run = paragraph.add_run(f" ({tag})")
            _set_run_font(run)
        return True
    element = _formula_omml(source, display=display)
    if element is None:
        fallback = _formula_fallback_text(source)
        if fallback:
            run = paragraph.add_run(fallback)
            _set_run_font(run)
            return bool(fallback)
        return False
    paragraph._p.append(element)
    _, tag = _formula_parts(source)
    if tag:
        run = paragraph.add_run(f" ({tag})")
        _set_run_font(run)
    return True


def _inline_fragment_text(value: object) -> str:
    raw = str(value or "")
    text = plain_markdown_text(raw)
    if not text:
        return " " if raw and raw.isspace() else ""
    if raw[:1].isspace():
        text = " " + text
    if raw[-1:].isspace():
        text += " "
    return text


def _set_run_font(run, name: str = "Calibri") -> None:
    has_cjk = bool(re.search(r"[\u2e80-\u9fff\uf900-\ufaff]", run.text or ""))
    resolved_name = CJK_FONT if has_cjk else name
    run.font.name = resolved_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), resolved_name)
    r_fonts.set(qn("w:hAnsi"), resolved_name)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    if has_cjk:
        r_fonts.set(qn("w:hint"), "eastAsia")
        language = r_pr.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            r_pr.append(language)
        language.set(qn("w:val"), "zh-CN")
        language.set(qn("w:eastAsia"), "zh-CN")


def _add_plain_inline_markdown(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(_inline_fragment_text(text[position : match.start()]))
            _set_run_font(run)
        token = match.group(0)
        run = paragraph.add_run(plain_markdown_text(token))
        _set_run_font(run)
        if token.startswith(("**", "__")):
            run.bold = True
        elif token.startswith("`"):
            _set_run_font(run, "Courier New")
            run.font.size = Pt(9.5)
        else:
            run.italic = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(_inline_fragment_text(text[position:]))
        _set_run_font(run)


def _add_inline_markdown(
    paragraph,
    text: str,
    *,
    display_math: bool = False,
    table_cell: bool = False,
) -> None:
    """Add Markdown text while turning delimited LaTeX into formula images."""
    position = 0
    source_text = str(text or "")
    for match in _iter_formula_matches(source_text):
        if match.start() > position:
            _add_plain_inline_markdown(paragraph, source_text[position : match.start()])
        source = _formula_source_from_match(match)
        is_display = _formula_match_is_display(match)
        _append_formula(
            paragraph,
            source,
            display=display_math and is_display,
            table_cell=table_cell,
        )
        position = match.end()
    if position < len(source_text):
        _add_plain_inline_markdown(paragraph, source_text[position:])


def _is_standalone_display_formula(text: str) -> bool:
    """Return whether *text* is a complete display-math block.

    Equation labels such as ``(5)`` are deliberately treated as mixed text;
    they are appended as a normal run after the OMML object.  This avoids
    placing a Word paragraph-level math object next to an ordinary suffix.
    """

    value = str(text or "").strip()
    matches = list(_iter_formula_matches(value))
    if len(matches) != 1:
        return False
    match = matches[0]
    if not _formula_match_is_display(match):
        return False
    return not value[: match.start()].strip() and not value[match.end() :].strip()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def _table_column_widths(rows: list[list[str]], total_width: int) -> list[int]:
    column_count = len(rows[0])
    weights: list[float] = []
    for column_index in range(column_count):
        values = [plain_markdown_text(row[column_index]) for row in rows if column_index < len(row)]
        non_empty = [value for value in values if value]
        numeric = bool(non_empty) and all(
            re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?%?", value.strip()) for value in non_empty[1:]
        )
        longest = max((max((len(line) for line in value.splitlines()), default=0) for value in non_empty), default=0)
        weights.append(6.0 if numeric else float(min(28, max(7, longest))))
    total_weight = sum(weights) or float(column_count)
    widths = [max(420, round(total_width * weight / total_weight)) for weight in weights]
    difference = total_width - sum(widths)
    widths[-1] += difference
    if widths[-1] < 420:
        deficit = 420 - widths[-1]
        widths[-1] = 420
        donor = max(range(max(1, column_count - 1)), key=lambda index: widths[index])
        widths[donor] = max(420, widths[donor] - deficit)
    return widths


def _configure_table_geometry(
    table,
    rows: list[list[str]],
    *,
    total_width: int = 9360,
) -> None:
    column_count = len(rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    column_widths = _table_column_widths(rows, total_width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in column_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(column_width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, column_widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def _configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    _add_page_number(section.footer.paragraphs[0])


def _set_section_layout(section, *, landscape: bool) -> None:
    """Apply a deterministic page layout to a newly-created section."""

    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _image_lookup_keys(task: dict, path: str) -> list[str]:
    """Resolve Markdown/HTML image names to the task's stored image keys."""

    requested = str(path or "").strip()
    normalized = requested.lstrip("./")
    candidates: list[str] = [requested, normalized]
    if normalized.startswith("imgs/"):
        candidates.append("ocr_images/" + normalized[5:])
    basename = Path(normalized).name

    # PaddleOCR-VL stores page images under ``ocr_images/<batch-id>_...`` but
    # emits ``imgs/<name>`` in the HTML.  The per-page markdown metadata keeps
    # the exact alias; consult it before falling back to a unique basename.
    for record in task.get("ocrResults") or []:
        markdown = record.get("markdown") if isinstance(record, dict) else None
        if isinstance(markdown, dict):
            mapped = (markdown.get("images") or {}).get(requested)
            if isinstance(mapped, str):
                candidates.append(mapped)
            for source, mapped in (markdown.get("images") or {}).items():
                if str(source).lstrip("./") == normalized and isinstance(mapped, str):
                    candidates.append(mapped)
    batch_markdown = task.get("batchMarkdown") or {}
    if isinstance(batch_markdown, dict):
        for value in batch_markdown.values():
            markdown = value.get("markdown") if isinstance(value, dict) else value
            if isinstance(markdown, dict):
                mapped = (markdown.get("images") or {}).get(requested)
                if isinstance(mapped, str):
                    candidates.append(mapped)

    images = task.get("images") or {}
    if basename:
        candidates.extend(
            key for key in images
            if isinstance(key, str) and Path(key).name == basename
        )
    # Preserve insertion order while avoiding repeated decodes.
    return list(dict.fromkeys(candidates))


def _image_bytes(task: dict, path: str) -> bytes | None:
    images = task.get("images") or {}
    payload = None
    for key in _image_lookup_keys(task, path):
        if key in images:
            payload = images.get(key)
            break
    if not isinstance(payload, str) or not payload:
        return None
    if payload.startswith("data:"):
        payload = payload.split(",", 1)[-1]
    try:
        data = base64.b64decode(payload, validate=True)
    except (ValueError, TypeError):
        return None
    if len(data) > 25 * 1024 * 1024:
        return None
    try:
        with Image.open(io.BytesIO(data)) as image:
            image.verify()
    except Exception:
        return None
    return data


def _add_markdown_image(
    document: Document,
    task: dict,
    alt: str,
    path: str,
    *,
    width_in: float | None = None,
) -> bool:
    data = _image_bytes(task, path)
    if data is None:
        return False
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(io.BytesIO(data), width=Inches(width_in or 6.25))
    except Exception:
        paragraph._element.getparent().remove(paragraph._element)
        return False
    if alt:
        caption = document.add_paragraph(plain_markdown_text(alt))
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for caption_run in caption.runs:
            caption_run.italic = True
            caption_run.font.size = Pt(9)
            caption_run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
    return True


def _parse_html_image_line(line: str) -> tuple[str, str, float | None] | None:
    """Extract ``src``, ``alt`` and an optional percentage width from OCR HTML."""

    match = HTML_IMAGE_LINE_RE.fullmatch(str(line or ""))
    if match is None:
        return None
    attrs: dict[str, str] = {}
    for attr in HTML_IMAGE_ATTR_RE.finditer(match.group("attrs") or ""):
        attrs[attr.group("name").lower()] = html.unescape(
            attr.group("double") or attr.group("single") or attr.group("bare") or ""
        )
    source = attrs.get("src", "").strip()
    if not source:
        return None
    alt = attrs.get("alt", "").strip()
    width_in: float | None = None
    width = attrs.get("width", "").strip()
    percent = re.fullmatch(r"(\d+(?:\.\d+)?)\s*%", width)
    pixels = re.fullmatch(r"(\d+(?:\.\d+)?)\s*(?:px)?", width, re.IGNORECASE)
    if percent:
        width_in = min(6.25, max(1.0, 6.25 * float(percent.group(1)) / 100.0))
    elif pixels:
        # PaddleOCR-VL's source canvas is normally 1191 px wide.  Convert its
        # pixel hint to the printable DOCX width while keeping a safe bound.
        width_in = min(6.25, max(1.0, 6.25 * float(pixels.group(1)) / 1191.0))
    return alt, source, width_in


def _add_docx_table(
    document: Document,
    rows: list[list[str]],
    *,
    merges: list[tuple[int, int, int, int]] | None = None,
    header_rows: int = 1,
    total_width: int = 9360,
) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    _configure_table_geometry(table, rows, total_width=total_width)
    # Parameter/key-value tables are conventionally centered in the value
    # column.  OCR often mixes numeric and textual values; aligning only the
    # numeric cells made one column look inconsistent in WPS/Word.
    header_text = [plain_markdown_text(str(v)).strip().lower() for v in rows[0]]
    value_column = (
        len(rows[0]) == 2
        and header_text[0] in {"name", "parameter", "key"}
        and header_text[1] in {"value", "setting", "result"}
    )
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            _add_inline_markdown(
                paragraph,
                str(value)
                .replace("<br>", "\n")
                .replace("<br/>", "\n")
                # PaddleOCR-VL serializes line breaks inside HTML table cells
                # as the two literal characters ``\\n``.  Turn those back
                # into real paragraph breaks before writing DOCX runs.
                .replace(r"\n", "\n"),
                table_cell=True,
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if row_index < header_rows:
                _set_cell_shading(cell, "E8EEF5")
                for run in paragraph.runs:
                    run.bold = True
            if (value_column and row_index > 0 and column_index == 1) or (
                value and re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?%?", plain_markdown_text(value).strip())
            ):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                if run.font.size is None:
                    run.font.size = Pt(7.5 if len(rows[0]) >= 8 else 8.5 if len(rows[0]) >= 5 else 9)
    for row_index in range(min(header_rows, len(table.rows))):
        _repeat_table_header(table.rows[row_index])
    for start_row, start_column, end_row, end_column in merges or []:
        if (
            0 <= start_row <= end_row < len(rows)
            and 0 <= start_column <= end_column < len(rows[0])
            and (start_row != end_row or start_column != end_column)
        ):
            table.cell(start_row, start_column).merge(table.cell(end_row, end_column))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_docx(task: dict) -> bytes:
    markdown, html_tables = _prepare_html_table_markers(str(task.get("markdown") or ""))
    if not markdown:
        raise ExportError("Task has no Markdown result")
    document = Document()
    _configure_document_styles(document)
    title = Path(str(task.get("name") or "OCR result")).stem
    document.core_properties.title = title
    document.core_properties.subject = "PaddleOCR Local editable export"
    document.core_properties.author = "PaddleOCR Local"
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    model = str(task.get("modelName") or task.get("modelId") or "PaddleOCR Local")
    subtitle = document.add_paragraph(f"{model} · PaddleOCR Local")
    subtitle.paragraph_format.space_after = Pt(12)
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)

    lines = markdown.split("\n")
    index = 0
    in_fence = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []
    # A wide HTML table is rendered in a landscape section.  Keep that
    # section active until the next actual non-wide content arrives instead
    # of immediately appending a portrait section after every table.  The
    # latter produced empty trailing pages for documents ending in a wide
    # table and unnecessary portrait transitions between adjacent wide
    # tables.
    landscape_section_open = False

    def ensure_portrait_section() -> None:
        nonlocal landscape_section_open
        if landscape_section_open:
            portrait_section = document.add_section(WD_SECTION.NEW_PAGE)
            _set_section_layout(portrait_section, landscape=False)
            landscape_section_open = False

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        content = " ".join(part.strip() for part in paragraph_lines)
        paragraph = document.add_paragraph()
        # A paragraph containing only a ``$$...$$``/``\[...\]`` block can use
        # the paragraph-level OMML container.  Mixed prose stays inline so the
        # surrounding text remains valid WordprocessingML.
        display_math = _is_standalone_display_formula(content)
        if display_math:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline_markdown(paragraph, content, display_math=display_math)
        paragraph_lines.clear()

    while index < len(lines):
        line = lines[index]
        if landscape_section_open and line.strip():
            next_marker = re.fullmatch(r"__PANDOCR_HTML_TABLE_(\d+)__", line.strip())
            is_next_wide_table = False
            if next_marker:
                next_index = int(next_marker.group(1))
                if 0 <= next_index < len(html_tables):
                    is_next_wide_table = len(html_tables[next_index].rows[0]) >= 9
            if not is_next_wide_table:
                ensure_portrait_section()
        if re.match(r"^\s*Algorithm\s+\d+\b", line, re.IGNORECASE):
            flush_paragraph()
            algorithm_lines = [line.strip()]
            index += 1
            saw_nonempty = False
            while index < len(lines):
                candidate = lines[index].strip()
                # HPD-Parsing emits the algorithm body as one long line and
                # terminates it with a blank line, unlike NaviDC's numbered
                # line-by-line form which runs through ``Output:``.
                if not candidate and saw_nonempty:
                    break
                algorithm_lines.append(candidate)
                saw_nonempty = saw_nonempty or bool(candidate)
                if re.match(r"^\s*Output\s*:", candidate, re.IGNORECASE):
                    index += 1
                    break
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.18)
            paragraph.paragraph_format.space_after = Pt(8)
            for line_index, algorithm_line in enumerate(algorithm_lines):
                if line_index:
                    paragraph.add_run().add_break()
                # Algorithm blocks frequently contain inline/display LaTeX.
                # Route each source line through the same Markdown+OMML path
                # as ordinary prose so commands such as ``\\array`` and
                # ``\\operatorname`` are not flattened into literal text.
                run_start = len(paragraph.runs)
                _add_inline_markdown(paragraph, algorithm_line)
                for run in paragraph.runs[run_start:]:
                    _set_run_font(run, "Courier New")
                    run.font.size = Pt(9)
            continue
        marker = re.fullmatch(r"__PANDOCR_HTML_TABLE_(\d+)__", line.strip())
        if marker:
            flush_paragraph()
            table_index = int(marker.group(1))
            if 0 <= table_index < len(html_tables):
                table_spec = html_tables[table_index]
                # Very wide scientific result tables need a landscape page;
                # keeping the rest of the report portrait preserves normal
                # reading flow and prevents method names from being split into
                # one character per line.
                if len(table_spec.rows[0]) >= 9:
                    if not landscape_section_open:
                        landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
                        _set_section_layout(landscape_section, landscape=True)
                        landscape_section_open = True
                    _add_docx_table(
                        document,
                        table_spec.rows,
                        merges=table_spec.merges,
                        header_rows=table_spec.header_rows,
                        total_width=14250,
                    )
                else:
                    ensure_portrait_section()
                    _add_docx_table(
                        document,
                        table_spec.rows,
                        merges=table_spec.merges,
                        header_rows=table_spec.header_rows,
                    )
            index += 1
            continue
        if FENCE_RE.match(line):
            flush_paragraph()
            if in_fence:
                paragraph = document.add_paragraph("\n".join(code_lines))
                paragraph.style = document.styles["Normal"]
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.right_indent = Inches(0.18)
                for run in paragraph.runs:
                    _set_run_font(run, "Courier New")
                    run.font.size = Pt(9)
                code_lines.clear()
            in_fence = not in_fence
            index += 1
            continue
        if in_fence:
            code_lines.append(line)
            index += 1
            continue
        if index + 1 < len(lines) and "|" in line and is_markdown_table_separator(lines[index + 1]):
            flush_paragraph()
            header = split_markdown_table_row(line)
            rows = [header]
            index += 2
            while index < len(lines) and "|" in lines[index] and not FENCE_RE.match(lines[index]):
                row = split_markdown_table_row(lines[index])
                if len(row) < 2:
                    break
                rows.append((row + [""] * len(header))[: len(header)])
                index += 1
            _add_docx_table(document, rows)
            continue
        if not line.strip():
            flush_paragraph()
            index += 1
            continue
        image_match = MARKDOWN_IMAGE_RE.fullmatch(line.strip())
        if image_match:
            flush_paragraph()
            if not _add_markdown_image(document, task, image_match.group(1), image_match.group(2)):
                paragraph_lines.append(image_match.group(1) or image_match.group(2))
            index += 1
            continue
        html_image = _parse_html_image_line(line)
        if html_image is not None:
            flush_paragraph()
            alt, source, width_in = html_image
            if not _add_markdown_image(document, task, alt, source, width_in=width_in):
                paragraph_lines.append(alt or source)
            index += 1
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)), 3)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline_markdown(paragraph, heading_match.group(2))
            index += 1
            continue
        list_match = re.match(r"^\s*([-*+]|\d+[.)])\s+(.+)$", line)
        if list_match:
            flush_paragraph()
            style = "List Number" if list_match.group(1)[0].isdigit() else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            _add_inline_markdown(paragraph, list_match.group(2))
            index += 1
            continue
        if line.lstrip().startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            _add_inline_markdown(paragraph, line.lstrip()[1:].strip())
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0x52, 0x60, 0x76)
                run.italic = True
            index += 1
            continue
        if re.fullmatch(r"\s*([-*_])\1{2,}\s*", line):
            flush_paragraph()
            index += 1
            continue
        paragraph_lines.append(line)
        index += 1
    if in_fence and code_lines:
        paragraph_lines.extend(code_lines)
    flush_paragraph()

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _safe_excel_value(value: object) -> str:
    text = plain_markdown_text(value).replace("<br>", "\n").replace("<br/>", "\n")
    if text.startswith(FORMULA_PREFIXES):
        return "'" + text
    return text


def build_xlsx(task: dict) -> bytes:
    markdown = str(task.get("markdown") or "")
    tables = extract_markdown_tables(markdown) + extract_html_tables(markdown)
    if not tables:
        raise ExportError("Task has no Markdown tables")
    workbook = Workbook()
    workbook.remove(workbook.active)
    header_fill = PatternFill("solid", fgColor="E8EEF5")
    header_font = Font(bold=True, color="0B2545")
    for table_index, rows in enumerate(tables, start=1):
        sheet = workbook.create_sheet(f"Table {table_index}")
        for row_index, values in enumerate(rows, start=1):
            for column_index, value in enumerate(values, start=1):
                cell = sheet.cell(row_index, column_index, _safe_excel_value(value))
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if row_index == 1:
                    cell.fill = header_fill
                    cell.font = header_font
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for column_index in range(1, len(rows[0]) + 1):
            values = [str(sheet.cell(row, column_index).value or "") for row in range(1, len(rows) + 1)]
            width = min(60, max(10, max((len(part) for value in values for part in value.split("\n")), default=10) + 2))
            sheet.column_dimensions[get_column_letter(column_index)].width = width
        sheet.sheet_view.showGridLines = False
    workbook.properties.title = Path(str(task.get("name") or "OCR tables")).stem
    workbook.properties.creator = "PaddleOCR Local"
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _source_as_pdf(source_path: Path, mime_type: str) -> bytes:
    content = source_path.read_bytes()
    if mime_type == "application/pdf" or content.startswith(b"%PDF"):
        try:
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted and not reader.decrypt(""):
                raise ExportError("Source PDF is encrypted")
        except Exception as error:
            if isinstance(error, ExportError):
                raise
            raise ExportError("Source PDF is invalid or encrypted") from error
        return content
    try:
        with Image.open(io.BytesIO(content)) as source_image:
            image = ImageOps.exif_transpose(source_image).convert("RGB")
            width, height = image.size
            max_width, max_height = A4
            scale = min(max_width / width, max_height / height)
            page_size = (max(72, width * scale), max(72, height * scale))
            output = io.BytesIO()
            pdf = canvas.Canvas(output, pagesize=page_size, pageCompression=1)
            pdf.drawImage(ImageReader(image), 0, 0, width=page_size[0], height=page_size[1])
            pdf.showPage()
            pdf.save()
            return output.getvalue()
    except Exception as error:
        raise ExportError("Source file is not a supported PDF or image") from error


def _page_number(result: dict, fallback: int) -> int:
    for key in ("sourcePage", "pageIndex", "page_index"):
        value = result.get(key)
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            page = int(value)
            return max(1, page if key == "sourcePage" else page + 1)
    return fallback


def _result_text_lines(result: dict) -> list[str]:
    lines: list[str] = []
    ocr_lines = result.get("ocrLines")
    if isinstance(ocr_lines, list):
        lines.extend(str(line.get("text") or "") for line in ocr_lines if isinstance(line, dict))
    blocks = result.get("parsing_res_list") or result.get("parsingResList")
    if isinstance(blocks, list):
        for block in blocks:
            if not isinstance(block, dict):
                continue
            value = block.get("block_content") or block.get("blockContent") or block.get("text")
            if value:
                lines.extend(str(value).splitlines())
    markdown = result.get("markdown")
    if isinstance(markdown, dict):
        markdown = markdown.get("text")
    if not lines and isinstance(markdown, str):
        lines.extend(markdown.splitlines())
    if not lines and isinstance(result.get("text"), str):
        lines.extend(result["text"].splitlines())
    output: list[str] = []
    seen: set[str] = set()
    for line in lines:
        clean = plain_markdown_text(line)
        if clean and clean not in seen:
            output.append(clean)
            seen.add(clean)
    return output


def _markdown_page_text(task: dict, page_count: int) -> list[list[str]]:
    pages: list[list[str]] = [[] for _ in range(page_count)]
    markdown = normalize_markdown(task.get("markdown"))
    if not markdown:
        return pages
    fragments = re.split(r"\n\s*---\s*\n", markdown)
    for index, fragment in enumerate(fragments[:page_count]):
        pages[index] = [plain_markdown_text(line) for line in fragment.splitlines() if plain_markdown_text(line)]
    if len(fragments) == 1:
        pages[0] = [plain_markdown_text(line) for line in markdown.splitlines() if plain_markdown_text(line)]
    return pages


def _task_page_text(task: dict, page_count: int) -> list[list[str]]:
    if task.get("manualEditedAt") and normalize_markdown(task.get("markdown")):
        return _markdown_page_text(task, page_count)
    pages: list[list[str]] = [[] for _ in range(page_count)]
    for fallback, result in enumerate(task.get("ocrResults") or [], start=1):
        if not isinstance(result, dict):
            continue
        page = min(page_count, _page_number(result, fallback))
        pages[page - 1].extend(_result_text_lines(result))
    if any(pages):
        return pages
    return _markdown_page_text(task, page_count)


def _register_search_font() -> str:
    name = "STSong-Light"
    if name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(name))
    return name


def _invisible_overlay(page_size: tuple[float, float], lines: Iterable[str]) -> bytes:
    width, height = page_size
    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=page_size, pageCompression=1)
    text = pdf.beginText(4, max(4, height - 8))
    text.setFont(_register_search_font(), 1)
    text.setLeading(1.2)
    text.setTextRenderMode(3)
    max_characters = max(80, int(width / 1.1))
    for value in lines:
        clean = "".join(character for character in str(value) if character >= " " or character == "\t")
        for offset in range(0, len(clean), max_characters):
            text.textLine(clean[offset : offset + max_characters])
            if text.getY() < 2:
                text.setTextOrigin(4, max(4, height - 8))
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()
    return output.getvalue()


def build_searchable_pdf(task: dict, source_path: Path) -> bytes:
    if not source_path.exists():
        raise ExportError("Task source file is missing")
    source_pdf = _source_as_pdf(source_path, str(task.get("mimeType") or ""))
    reader = PdfReader(io.BytesIO(source_pdf))
    if reader.is_encrypted:
        try:
            decrypted = reader.decrypt("")
        except Exception as error:
            raise ExportError("Source PDF is encrypted") from error
        if not decrypted:
            raise ExportError("Source PDF is encrypted")
    pages = _task_page_text(task, len(reader.pages))
    if not any(pages):
        raise ExportError("Task has no OCR text to add")
    writer = PdfWriter()
    for index, page in enumerate(reader.pages):
        writer.add_page(page)
        output_page = writer.pages[-1]
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        if pages[index]:
            overlay = PdfReader(io.BytesIO(_invisible_overlay((width, height), pages[index]))).pages[0]
            output_page.merge_page(overlay, over=True)
    metadata = dict(reader.metadata or {})
    metadata["/Producer"] = "PaddleOCR Local"
    metadata["/Subject"] = "Searchable offline OCR export"
    writer.add_metadata({str(key): str(value) for key, value in metadata.items() if value is not None})
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()
